import asyncio
import contextlib
import io
import logging
import secrets
import socket
from pathlib import Path

import anthropic
import httpx
from fastapi import FastAPI, Request, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

STATIC_DIR = Path(__file__).resolve().parent / "static"
MAX_UPLOAD_BYTES = 10 * 1024 * 1024

log = logging.getLogger("app.server")


def _extract_text(filename: str, data: bytes) -> str:
    """Pull plain text out of an uploaded resume/notes file."""
    suffix = Path(filename or "").suffix.lower()
    if suffix in (".txt", ".md", ""):
        return data.decode("utf-8", errors="replace")
    if suffix == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(filter(None, (page.extract_text() for page in reader.pages)))
    if suffix == ".docx":
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    raise ValueError(f"Unsupported file type '{suffix or filename}'. Use .txt, .md, .pdf, or .docx.")


def find_free_port(start: int = 8756, attempts: int = 50) -> int:
    for port in range(start, start + attempts):
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            try:
                s.bind(("127.0.0.1", port))
                return port
            except OSError:
                continue
    raise RuntimeError("No free port available for the UI server.")


async def _test_deepgram_key(key: str) -> str | None:
    """Return None if the key works, else a plain-English problem."""
    if not key:
        return "Paste your Deepgram key first."
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(
                "https://api.deepgram.com/v1/projects",
                headers={"Authorization": f"Token {key}"},
            )
        if resp.status_code == 200:
            return None
        if resp.status_code in (401, 403):
            return "Deepgram rejected this key — copy it again from console.deepgram.com."
        return f"Deepgram answered with an unexpected error ({resp.status_code})."
    except httpx.HTTPError:
        return "Could not reach Deepgram — check your internet connection."


async def _test_anthropic_key(key: str) -> str | None:
    if not key:
        return "Paste your Anthropic key first."

    def check() -> str | None:
        try:
            client = anthropic.Anthropic(api_key=key, timeout=10.0, max_retries=0)
            client.models.list(limit=1)  # free call; fails fast on a bad key
            return None
        except anthropic.AuthenticationError:
            return "Anthropic rejected this key — copy it again from console.anthropic.com."
        except anthropic.APIConnectionError:
            return "Could not reach Anthropic — check your internet connection."
        except anthropic.APIStatusError as e:
            return f"Anthropic answered with an unexpected error ({e.status_code})."
        except Exception as e:
            return f"Could not check the key: {e}"

    return await asyncio.to_thread(check)


def create_app(engine, port: int) -> FastAPI:
    # Random per-run token: pages served by us embed it; websocket and API
    # calls must present it. Combined with the Origin/Host checks below this
    # keeps other websites and other machines out of the session (the browser
    # does NOT enforce same-origin for WebSockets or localhost requests).
    token = secrets.token_urlsafe(24)
    allowed_hosts = {f"127.0.0.1:{port}", f"localhost:{port}"}
    allowed_origins = {f"http://{h}" for h in allowed_hosts}

    clients: set[WebSocket] = set()
    outbox: asyncio.Queue = asyncio.Queue()
    loop_holder: dict = {}
    broadcaster_task: dict = {}

    def emit(event: dict):
        """Called from engine threads; hands the event to the asyncio loop."""
        loop = loop_holder.get("loop")
        if loop is None:
            return
        try:
            loop.call_soon_threadsafe(outbox.put_nowait, event)
        except RuntimeError:
            pass

    async def broadcaster():
        while True:
            event = await outbox.get()
            for ws in list(clients):
                try:
                    await ws.send_json(event)
                except Exception:
                    clients.discard(ws)

    @contextlib.asynccontextmanager
    async def lifespan(app: FastAPI):
        loop_holder["loop"] = asyncio.get_running_loop()
        broadcaster_task["task"] = asyncio.create_task(broadcaster())
        on_ready = getattr(app.state, "on_ready", None)
        if on_ready:
            on_ready()
        yield
        task = broadcaster_task.get("task")
        if task:
            task.cancel()

    app = FastAPI(title="Interview Copilot", lifespan=lifespan)
    app.state.emit = emit
    app.state.token = token

    @app.middleware("http")
    async def guard(request: Request, call_next):
        # DNS-rebinding protection: only our own loopback hostnames.
        if request.headers.get("host", "") not in allowed_hosts:
            return JSONResponse({"error": "Forbidden."}, status_code=403)
        # State-changing endpoints additionally need the session token.
        if request.url.path.startswith("/api/"):
            if request.headers.get("x-auth-token", "") != token:
                return JSONResponse({"error": "Forbidden."}, status_code=403)
        return await call_next(request)

    @app.get("/")
    async def index():
        html = (STATIC_DIR / "index.html").read_text(encoding="utf-8")
        return HTMLResponse(html.replace("__AUTH_TOKEN__", token))

    app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

    # -- setup wizard --------------------------------------------------------

    @app.post("/api/setup/test")
    async def setup_test(payload: dict):
        service = str(payload.get("service", ""))
        key = str(payload.get("key", "")).strip()
        if service == "deepgram":
            problem = await _test_deepgram_key(key)
        elif service == "anthropic":
            problem = await _test_anthropic_key(key)
        else:
            problem = "Unknown service."
        return {"ok": problem is None, "error": problem}

    @app.post("/api/setup/keys")
    async def setup_keys(payload: dict):
        dg = str(payload.get("deepgram", "")).strip()
        ant = str(payload.get("anthropic", "")).strip()
        for name, problem in (
            ("deepgram", await _test_deepgram_key(dg)),
            ("anthropic", await _test_anthropic_key(ant)),
        ):
            if problem:
                return JSONResponse({"ok": False, "field": name, "error": problem}, status_code=400)
        try:
            await asyncio.to_thread(engine.configure_keys, dg, ant)
        except Exception as e:
            log.exception("configure_keys failed")
            return JSONResponse(
                {"ok": False, "error": f"Could not finish setup: {e}"}, status_code=500,
            )
        return {"ok": True}

    # -- file extraction (profile uploads) -----------------------------------

    @app.post("/api/extract")
    async def extract(file: UploadFile):
        data = await file.read()
        if len(data) > MAX_UPLOAD_BYTES:
            return JSONResponse({"error": "File is too large (10 MB max)."}, status_code=413)
        try:
            text = _extract_text(file.filename or "", data)
        except ValueError as e:
            return JSONResponse({"error": str(e)}, status_code=400)
        except Exception as e:
            return JSONResponse({"error": f"Could not read the file: {e}"}, status_code=422)
        if not text.strip():
            return JSONResponse(
                {"error": "No text could be extracted (the file may be a scanned image)."},
                status_code=422,
            )
        return {"text": text.strip(), "filename": file.filename}

    # -- live event stream ---------------------------------------------------

    @app.websocket("/ws")
    async def ws_endpoint(ws: WebSocket):
        origin = ws.headers.get("origin", "")
        if origin not in allowed_origins or ws.query_params.get("token") != token:
            await ws.close(code=4403)
            return
        await ws.accept()
        clients.add(ws)
        try:
            for event in engine.snapshot():
                await ws.send_json(event)
            while True:
                command = await ws.receive_json()
                if isinstance(command, dict):
                    engine.submit(command)
        except WebSocketDisconnect:
            pass
        except Exception:
            pass
        finally:
            clients.discard(ws)

    return app
